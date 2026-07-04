# -*- coding: utf-8 -*-
"""解析英语题库（听力/阅读/补充练习/翻译），输出 questions.js"""
import docx
import json
import re
import os
from datetime import datetime

CJK_RE = re.compile(r'[\u4e00-\u9fff]')
ANS_RE = re.compile(r'正确答案为[：:]\s*([A-Z])')

# ---------- 听力 ----------
def parse_listening(filepath, source):
    doc = docx.Document(filepath)
    paras = [p.text.strip().replace('\xa0', ' ') for p in doc.paragraphs]
    questions = []
    i, n = 0, len(paras)
    SKIP = ('Section', 'Directions', '本题得分', '本小题得分', '查看基本信息', 'Script', '*')

    while i < n:
        line = paras[i]
        if not line or any(line.startswith(s) for s in SKIP):
            i += 1; continue

        # 题目行：1.* 或 1.（带*或不带）
        mq = re.match(r'^(\d+)\s*\.\s*\*?\s*(.*)$', line)
        # 选项行（14.docx 第一题缺题干，直接是A)）
        mo_opt = re.match(r'^([A-D])\s*\)', line)

        if mq or mo_opt:
            stem = ''
            if mq:
                stem = mq.group(2).strip()
                i += 1
            # 收集选项
            options = {}
            answer = []
            while i < n:
                nxt = paras[i]
                if not nxt:
                    i += 1; continue
                if any(nxt.startswith(s) for s in SKIP):
                    i += 1; continue
                mo = re.match(r'^([A-D])\s*\)\s*(.*)$', nxt)
                if mo:
                    L = mo.group(1)
                    options[L] = mo.group(2).strip()
                    i += 1
                    continue
                ma = ANS_RE.search(nxt)
                if ma:
                    answer = [ma.group(1)]
                    i += 1
                    break
                # 答案行可能跨行（"正确答案为： A  你错选为..."）
                if '正确答案' in nxt:
                    m2 = re.search(r'([A-D])', nxt)
                    if m2: answer = [m2.group(1)]
                    i += 1
                    break
                # 遇到下一题
                if re.match(r'^\d+\s*\.\s*\*?', nxt):
                    break
                i += 1

            if options:
                questions.append({
                    'stem': stem or '(本题无题干)',
                    'options': options,
                    'answer': answer,
                    'source': source
                })
            continue
        i += 1
    return questions

# ---------- 阅读 ----------
def parse_reading(filepath):
    doc = docx.Document(filepath)
    paras = [p.text.strip().replace('\xa0', ' ') for p in doc.paragraphs]
    questions = []
    i, n = 0, len(paras)
    passage = ''

    while i < n:
        line = paras[i]
        if not line:
            i += 1; continue
        # 篇章标记
        if re.match(r'^第[一二三四五六七八九十]+篇', line):
            passage = line
            i += 1; continue
        # 题目行：1.What...
        mq = re.match(r'^(\d+)\s*\.\s*(.*)$', line)
        if mq:
            stem = mq.group(2).strip()
            i += 1
            options = {}
            answer = []
            while i < n:
                nxt = paras[i]
                if not nxt:
                    i += 1; continue
                mo = re.match(r'^([A-D])\s*\)\s*(.*)$', nxt)
                if mo:
                    L = mo.group(1)
                    txt = mo.group(2).strip()
                    # 含中文 = 正确答案
                    if CJK_RE.search(txt):
                        answer = [L]
                        # 去掉中文部分用于显示
                        txt = CJK_RE.split(txt)[0].strip()
                    options[L] = txt
                    i += 1
                    continue
                # 遇到下一题或篇章
                if re.match(r'^\d+\s*\.', nxt) or re.match(r'^第[一二三四五六七八九十]+篇', nxt):
                    break
                # 跳过续行中文
                i += 1
            if options and answer:
                full_stem = (passage + ' ' if passage else '') + stem
                questions.append({
                    'stem': full_stem,
                    'options': options,
                    'answer': answer,
                    'source': os.path.basename(filepath)
                })
            continue
        i += 1
    return questions

# ---------- 补充练习：事实与观点 ----------
def parse_fo(filepath):
    doc = docx.Document(filepath)
    paras = [p.text.strip().replace('\xa0', ' ') for p in doc.paragraphs]
    questions = []
    i, n = 0, len(paras)
    while i < n:
        line = paras[i]
        if not line:
            i += 1; continue
        mq = re.match(r'^(\d+)\s*\.\s*(.*)$', line)
        if mq and not line.startswith('Direction'):
            stem = mq.group(2).strip()
            # 去掉分值标记 （6.5分）
            stem = re.sub(r'（\d+\.?\d*分）', '', stem).strip()
            i += 1
            # 找选项行（A. F    B. O   C. F / O）
            options = {}
            answer = []
            while i < n:
                nxt = paras[i]
                if not nxt:
                    i += 1; continue
                # 答案行
                if nxt.startswith('答案'):
                    m = re.search(r'([A-C])', nxt)
                    if m: answer = [m.group(1)]
                    i += 1
                    break
                # 选项行：A. F    B. O   C. F / O
                if re.match(r'^[A-C]\s*\.', nxt):
                    parts = re.split(r'([A-C])\s*\.', nxt)
                    for k in range(1, len(parts)-1, 2):
                        L = parts[k]
                        txt = parts[k+1].strip()
                        options[L] = txt
                    i += 1
                    continue
                # 解析行 - 跳过
                if nxt.startswith('解析'):
                    i += 1
                    continue
                # 下一题
                if re.match(r'^\d+\s*\.', nxt):
                    break
                i += 1
            if options and answer:
                questions.append({
                    'stem': stem,
                    'options': options,
                    'answer': answer,
                    'source': os.path.basename(filepath)
                })
            continue
        i += 1
    return questions

# ---------- 补充练习：名词化 ----------
def parse_nominalization(filepath):
    doc = docx.Document(filepath)
    paras = [p.text.strip().replace('\xa0', ' ') for p in doc.paragraphs]
    questions = []
    i, n = 0, len(paras)
    while i < n:
        line = paras[i]
        if not line:
            i += 1; continue
        # 题目：1. Group 1
        mq = re.match(r'^(\d+)\s*\.\s*(.*)$', line)
        if mq and 'Group' in line:
            group_label = line.strip()
            i += 1
            # 收集 1) 和 2) 句子
            sents = {}
            while i < n:
                nxt = paras[i]
                if not nxt:
                    i += 1; continue
                ms = re.match(r'^([12])\s*\)\s*(.*)$', nxt)
                if ms:
                    sents[ms.group(1)] = ms.group(2).strip()
                    i += 1
                    continue
                # 选项行 A. 1)   B. 2)
                if re.match(r'^[AB]\s*\.', nxt):
                    break
                if re.match(r'^\d+\s*\.', nxt):
                    break
                i += 1
            # 解析选项
            options = {}
            answer = []
            while i < n:
                nxt = paras[i]
                if not nxt:
                    i += 1; continue
                if nxt.startswith('答案'):
                    m = re.search(r'([AB])', nxt)
                    if m: answer = [m.group(1)]
                    i += 1
                    break
                if re.match(r'^[AB]\s*\.', nxt):
                    parts = re.split(r'([AB])\s*\.', nxt)
                    for k in range(1, len(parts)-1, 2):
                        L = parts[k]
                        num = parts[k+1].strip()
                        # 选项文本包含完整句子
                        txt = num + ' ' + sents.get(num, '')
                        options[L] = txt
                    i += 1
                    continue
                if nxt.startswith('解析'):
                    i += 1; continue
                if re.match(r'^\d+\s*\.', nxt):
                    break
                i += 1
            if options and answer:
                stem = 'Identify which sentence uses nominalization.\n' + group_label
                if '1' in sents: stem += '\n1) ' + sents['1']
                if '2' in sents: stem += '\n2) ' + sents['2']
                questions.append({
                    'stem': stem,
                    'options': options,
                    'answer': answer,
                    'source': os.path.basename(filepath)
                })
            continue
        i += 1
    return questions

# ---------- 翻译 ----------
def parse_translation(filepath):
    doc = docx.Document(filepath)
    paras = [p.text.strip().replace('\xa0', ' ') for p in doc.paragraphs]
    units = []
    cur_unit = None
    cur_pairs = []
    last_en = None
    for line in paras:
        if not line: continue
        if re.match(r'^Unit\s+\d+', line):
            if cur_unit and cur_pairs:
                units.append({'unit': cur_unit, 'pairs': cur_pairs})
            cur_unit = line
            cur_pairs = []
            last_en = None
            continue
        if cur_unit:
            if CJK_RE.search(line):
                # 中文段落
                if last_en:
                    cur_pairs.append({'en': last_en, 'zh': line})
                    last_en = None
                else:
                    cur_pairs.append({'en': '', 'zh': line})
            else:
                # 英文段落
                if last_en:
                    cur_pairs.append({'en': last_en, 'zh': ''})
                last_en = line
    if cur_unit and cur_pairs:
        units.append({'unit': cur_unit, 'pairs': cur_pairs})
    if last_en:
        # dangling EN
        if units:
            units[-1]['pairs'].append({'en': last_en, 'zh': ''})
    return units

# ========== 主流程 ==========
all_questions = []
qid = 0

# 听力
listening_files = [
    r"D:\qq\英语\听力\11.docx",
    r"D:\qq\英语\听力\12.docx",
    r"D:\qq\英语\听力\13.docx",
    r"D:\qq\英语\听力\14.docx",
]
for f in listening_files:
    src = os.path.basename(f)
    for q in parse_listening(f, src):
        qid += 1
        # 听力/阅读全是单选；无答案的题也标 single，UI 不判分
        tp = 'single' if len(q['answer'])<=1 else 'multi'
        all_questions.append({
            'id': qid, 'subject': '英语', 'chapter': '听力',
            'type': tp,
            **q
        })

# 阅读
for q in parse_reading(r"D:\qq\英语\阅读\英语阅读docx.docx"):
    qid += 1
    all_questions.append({
        'id': qid, 'subject': '英语', 'chapter': '阅读',
        'type': 'single' if len(q['answer'])==1 else 'multi',
        **q
    })

# 补充练习-事实与观点
fo_path = r"D:\qq\英语\补充练习\Unit 5（F or and O）补充练习.docx"
if os.path.exists(fo_path):
    for q in parse_fo(fo_path):
        qid += 1
        all_questions.append({
            'id': qid, 'subject': '英语', 'chapter': '事实与观点',
            'type': 'single' if len(q['answer'])==1 else 'multi',
            **q
        })

# 补充练习-名词化
nom_path = r"D:\qq\英语\补充练习\nominalization（Unit 6）.docx"
if os.path.exists(nom_path):
    for q in parse_nominalization(nom_path):
        qid += 1
        all_questions.append({
            'id': qid, 'subject': '英语', 'chapter': '名词化',
            'type': 'single' if len(q['answer'])==1 else 'multi',
            **q
        })

# 翻译（非题目，单独存储）
translations = parse_translation(r"D:\qq\英语\翻译\翻译.docx")

# 输出
version = datetime.now().strftime("%Y%m%d%H%M%S")
js_out = r"e:\网课wqdwdwdqwd\questions.js"
with open(js_out, 'w', encoding='utf-8') as fp:
    fp.write('window.QUESTIONS=' + json.dumps(all_questions, ensure_ascii=False) + ';\n')
    fp.write('window.TRANSLATIONS=' + json.dumps(translations, ensure_ascii=False) + ';\n')
    fp.write('window.QUESTIONS_VERSION="' + version + '";\n')

# 统计
from collections import Counter
print(f'总题数: {len(all_questions)} | 版本: {version}')
ch = Counter(q['chapter'] for q in all_questions)
print('各板块:')
for c, v in ch.items():
    types = Counter(q['type'] for q in all_questions if q['chapter']==c)
    print(f'  {c}: {v}题 (单选{types["single"]}/多选{types["multi"]})')
no_ans = sum(1 for q in all_questions if not q['answer'])
print(f'无答案: {no_ans}')
print(f'翻译单元: {len(translations)} (共{sum(len(u["pairs"]) for u in translations)}段)')
